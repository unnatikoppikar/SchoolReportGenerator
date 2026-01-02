"""
Create sample test files for Report Card Generator.
Run this to generate sample Excel, Word template, and mapping files.
"""

import sys
from pathlib import Path

# Add parent to path
sys.path.insert(0, str(Path(__file__).parent.parent))

import pandas as pd
from docx import Document


def create_sample_excel(output_path: str):
    """Create a sample Excel file with student data."""
    print(f"Creating sample Excel: {output_path}")
    
    # Create DataFrame with 4 header rows + student data
    data = {
        0: ["GREENWOOD PUBLIC SCHOOL", "Class: 5A", "Academic Year: 2025-26", "Roll No", 
            "1", "2", "3", "4", "5"],
        1: ["", "Class Teacher: Mrs. Sharma", "", "Name",
            "Rahul Kumar", "Priya Singh", "Amit Verma", "Sneha Gupta", "Vikram Patel"],
        2: ["", "", "", "English",
            "85", "92", "45", "78", "65"],
        3: ["", "", "", "Hindi",
            "78", "88", "52", "82", "70"],
        4: ["", "", "", "Maths",
            "92", "95", "38", "75", "68"],
        5: ["", "", "", "Science",
            "88", "91", "48", "80", "72"],
        6: ["", "", "", "Total",
            "343", "366", "183", "315", "275"],
        7: ["", "", "", "Percentage",
            "85.75%", "91.50%", "45.75%", "78.75%", "68.75%"],
        8: ["", "", "", "Result",
            "Pass", "Pass", "Fail", "Pass", "Pass"],
        9: ["", "", "", "Grade",
            "A", "A+", "D", "B+", "B"]
    }
    
    df = pd.DataFrame(data)
    df.to_excel(output_path, index=False, header=False)
    print(f"  ✓ Created: {output_path}")


def create_sample_template(output_path: str):
    """Create a sample Word template with placeholders."""
    print(f"Creating sample Word template: {output_path}")
    
    doc = Document()
    
    # Title
    title = doc.add_paragraph()
    title.alignment = 1  # Center
    run = title.add_run("GREENWOOD PUBLIC SCHOOL")
    run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = 1
    run = subtitle.add_run("REPORT CARD")
    run.bold = True
    
    doc.add_paragraph()
    
    # Student info
    doc.add_paragraph("Student Name: {{name}}")
    doc.add_paragraph("Roll No: {{rollno}}")
    doc.add_paragraph("Class: {{class}}")
    
    doc.add_paragraph()
    
    # Marks section
    doc.add_paragraph("MARKS").runs[0].bold = True
    doc.add_paragraph("English: {{english}}")
    doc.add_paragraph("Hindi: {{hindi}}")
    doc.add_paragraph("Mathematics: {{maths}}")
    doc.add_paragraph("Science: {{science}}")
    
    doc.add_paragraph()
    
    # Summary
    doc.add_paragraph("SUMMARY").runs[0].bold = True
    doc.add_paragraph("Total Marks: {{total}}")
    doc.add_paragraph("Percentage: {{percentage}}")
    doc.add_paragraph("Result: {{result}}")
    doc.add_paragraph("Grade: {{grade}}")
    
    doc.save(output_path)
    print(f"  ✓ Created: {output_path}")


def create_sample_mapping(output_path: str):
    """Create a sample mapping JSON file."""
    print(f"Creating sample mapping: {output_path}")
    
    import json
    mapping = {
        "rollno": "A",
        "name": "B",
        "english": "C",
        "hindi": "D",
        "maths": "E",
        "science": "F",
        "total": "G",
        "percentage": "H",
        "result": "I",
        "grade": "J"
    }
    
    with open(output_path, 'w') as f:
        json.dump(mapping, f, indent=4)
    
    print(f"  ✓ Created: {output_path}")


def main():
    # Output directory
    output_dir = Path(__file__).parent.parent / "input_files"
    output_dir.mkdir(exist_ok=True)
    
    print("\n" + "="*50)
    print("Creating Sample Files")
    print("="*50 + "\n")
    
    # Create files
    create_sample_excel(str(output_dir / "sample_students.xlsx"))
    
    # Import docx here to handle missing module gracefully
    try:
        import docx
        create_sample_template(str(output_dir / "sample_template.docx"))
    except ImportError:
        print("  ⚠ python-docx not installed, skipping Word template")
    
    create_sample_mapping(str(output_dir / "sample_mapping.json"))
    
    print("\n" + "="*50)
    print("✓ Sample files created in: " + str(output_dir))
    print("="*50 + "\n")


if __name__ == "__main__":
    main()

